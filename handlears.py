import asyncio
import os
import re
import tempfile
from contextlib import suppress
from datetime import datetime

from telegram import Update
from telegram.error import BadRequest
from telegram.constants import ParseMode
from telegram.ext import (
    Application,
    CommandHandler,
    ContextTypes,
    ConversationHandler,
    MessageHandler,
    filters,
)

from config import BotConfig
from gemini_service import GeminiServiceError, build_structured_note, transcribe_audio
from keyboards import (
    BTN_CREATE_ROUTING,
    BTN_HELP,
    BTN_HISTORY,
    BTN_LANGUAGE,
    home_keyboard,
)

WAIT_ROUTING_INPUT = 1
TELEGRAM_TEXT_LIMIT = 3500


async def start_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    del context
    if not update.message:
        return

    await update.message.reply_text(
        text='Assalomu alaykum! Kerakli bo\'limni tanlang:',
        reply_markup=home_keyboard(),
    )


async def routing_entry_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    del context
    if update.message:
        await update.message.reply_text(
            'Audio yoki matn yuboring. Faqat 1 ta xabar qabul qilinadi.'
        )
    return WAIT_ROUTING_INPUT


async def routing_input_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    if not update.message:
        return ConversationHandler.END

    config: BotConfig = context.application.bot_data['config']

    processing_msg = await update.message.reply_text('Qayta ishlash jarayonida...')

    temp_audio_path: str | None = None
    try:
        source_text = ''

        if update.message.voice or update.message.audio or (
            update.message.document and update.message.document.mime_type and update.message.document.mime_type.startswith('audio/')
        ):
            temp_audio_path = await _download_audio_to_temp(update, context)
            source_text = await asyncio.to_thread(transcribe_audio, config.gemini_api_key, temp_audio_path)
        elif update.message.text:
            source_text = update.message.text.strip()
        else:
            await update.message.reply_text('Faqat audio yoki matn yuborishingiz mumkin.')
            return ConversationHandler.END

        result_text = await asyncio.to_thread(build_structured_note, config.gemini_api_key, source_text)

        with suppress(BadRequest):
            await processing_msg.delete()

        cleaned_text = _normalize_plain_medical_text(result_text)
        for chunk in _split_text_for_telegram(cleaned_text, TELEGRAM_TEXT_LIMIT):
            await update.message.reply_text(f'```text\n{chunk}\n```', parse_mode=ParseMode.MARKDOWN)

        docx_path = await asyncio.to_thread(_build_docx_from_markdown, cleaned_text)
        try:
            with open(docx_path, 'rb') as doc_file:
                await update.message.reply_document(
                    document=doc_file,
                    filename=os.path.basename(docx_path),
                    caption='WORD format tayyorlandi.',
                )
        finally:
            if os.path.exists(docx_path):
                with suppress(OSError):
                    os.remove(docx_path)
    except GeminiServiceError as exc:
        with suppress(BadRequest):
            await processing_msg.delete()
        await update.message.reply_text(f'Gemini xatosi: {exc}')
    except Exception as exc:  # noqa: BLE001
        with suppress(BadRequest):
            await processing_msg.delete()
        await update.message.reply_text(f'Xatolik yuz berdi: {exc}')
    finally:
        if temp_audio_path and os.path.exists(temp_audio_path):
            with suppress(OSError):
                os.remove(temp_audio_path)

    return ConversationHandler.END


async def help_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    del context
    if update.message:
        await update.message.reply_text('Yordam bo\'limi hozircha tayyorlanmoqda.')


async def history_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    del context
    if update.message:
        await update.message.reply_text('Tarix bo\'limi hozircha tayyorlanmoqda.')


async def language_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    del context
    if update.message:
        await update.message.reply_text('Tilni sozlash bo\'limi hozircha tayyorlanmoqda.')


async def _download_audio_to_temp(update: Update, context: ContextTypes.DEFAULT_TYPE) -> str:
    if not update.message:
        raise RuntimeError('Message not found')

    suffix = '.ogg'
    file_id = ''

    if update.message.voice:
        file_id = update.message.voice.file_id
        suffix = '.ogg'
    elif update.message.audio:
        file_id = update.message.audio.file_id
        if update.message.audio.file_name and '.' in update.message.audio.file_name:
            suffix = '.' + update.message.audio.file_name.rsplit('.', 1)[1]
    elif update.message.document and update.message.document.mime_type and update.message.document.mime_type.startswith('audio/'):
        file_id = update.message.document.file_id
        if update.message.document.file_name and '.' in update.message.document.file_name:
            suffix = '.' + update.message.document.file_name.rsplit('.', 1)[1]

    if not file_id:
        raise RuntimeError('Audio file_id topilmadi')

    fd, temp_path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)

    tg_file = await context.bot.get_file(file_id)
    await tg_file.download_to_drive(custom_path=temp_path)

    return temp_path


def _build_docx_from_markdown(markdown_text: str) -> str:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    cleaned = markdown_text.replace('```text', '').replace('```', '').strip()
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    previous_was_blank = False
    section_pattern = re.compile(r'^(\d+\.\s+[A-ZА-ЯЁʻ\'\-\s]+):\s*(.+)$')
    for raw_line in cleaned.splitlines():
        line = raw_line.strip()
        if not line:
            if not previous_was_blank:
                document.add_paragraph('')
            previous_was_blank = True
            continue

        previous_was_blank = False
        section_match = section_pattern.match(line)
        if section_match:
            heading_text = section_match.group(1)
            body_text = section_match.group(2)
            heading_paragraph = document.add_paragraph()
            heading_run = heading_paragraph.add_run(heading_text)
            heading_run.bold = True
            heading_run.font.name = 'Times New Roman'
            heading_run.font.size = Pt(14)

            body_paragraph = document.add_paragraph()
            body_run = body_paragraph.add_run(body_text)
            body_run.font.name = 'Times New Roman'
            body_run.font.size = Pt(14)
            continue

        paragraph = document.add_paragraph()
        tokens = re.split(r'(\*\*[^*]+\*\*)', line)
        for token in tokens:
            if not token:
                continue
            if token.startswith('**') and token.endswith('**') and len(token) > 4:
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            else:
                run = paragraph.add_run(token)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    file_path = os.path.join(tempfile.gettempdir(), f'tibbiy_hulosa_{timestamp}.docx')
    document.save(file_path)
    return file_path


def _normalize_plain_medical_text(text: str) -> str:
    cleaned_lines: list[str] = []
    for line in text.replace('```text', '').replace('```', '').splitlines():
        normalized = line.replace('**', '').replace('#', '').strip()
        if normalized:
            cleaned_lines.append(normalized)
    return '\n'.join(cleaned_lines)


def _split_text_for_telegram(text: str, limit: int) -> list[str]:
    if len(text) <= limit:
        return [text]

    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    for line in text.splitlines():
        extra = len(line) + (1 if current else 0)
        if current and current_len + extra > limit:
            chunks.append('\n'.join(current))
            current = [line]
            current_len = len(line)
        else:
            current.append(line)
            current_len += extra

    if current:
        chunks.append('\n'.join(current))

    normalized_chunks: list[str] = []
    for chunk in chunks:
        if len(chunk) <= limit:
            normalized_chunks.append(chunk)
            continue
        for i in range(0, len(chunk), limit):
            normalized_chunks.append(chunk[i:i + limit])
    return normalized_chunks


def register_handlers(application: Application, config: BotConfig) -> None:
    application.bot_data['config'] = config

    application.add_handler(CommandHandler('start', start_handler))

    text_filter = filters.TEXT & ~filters.COMMAND
    audio_or_text_filter = (
        filters.TEXT
        | filters.VOICE
        | filters.AUDIO
        | filters.Document.AUDIO
    )

    routing_conversation = ConversationHandler(
        entry_points=[
            MessageHandler(
                text_filter & filters.Regex(f'^{re.escape(BTN_CREATE_ROUTING)}$'),
                routing_entry_handler,
            )
        ],
        states={
            WAIT_ROUTING_INPUT: [
                MessageHandler(audio_or_text_filter, routing_input_handler),
            ]
        },
        fallbacks=[CommandHandler('start', start_handler)],
        allow_reentry=True,
    )

    application.add_handler(routing_conversation)
    application.add_handler(
        MessageHandler(text_filter & filters.Regex(f'^{re.escape(BTN_HELP)}$'), help_handler)
    )
    application.add_handler(
        MessageHandler(text_filter & filters.Regex(f'^{re.escape(BTN_HISTORY)}$'), history_handler)
    )
    application.add_handler(
        MessageHandler(text_filter & filters.Regex(f'^{re.escape(BTN_LANGUAGE)}$'), language_handler)
    )
